"""
Motor de plantillas con variables @.

Un modelo de escrito tiene texto con variables tipo @numero, @defendido, etc.
Al "armar" el escrito para un expediente, cada @variable se reemplaza por el
dato real. Si el dato falta (ej. no se cargó el defendido), queda un hueco
visible "[completar: ...]" en vez de tirar error.
"""

import re
from datetime import date

# ── Catálogo de variables disponibles ──────────────────────────
# (token, etiqueta legible, grupo) — es la fuente única que también ve el
# frontend para mostrar la ayuda. "Pasarse de opciones": mejor de más.
CATALOGO = [
    # Del expediente
    ("numero", "Número de expediente", "Expediente"),
    ("caratula", "Carátula", "Expediente"),
    ("juzgado", "Juzgado", "Expediente"),
    ("tipo_proceso", "Tipo de proceso", "Expediente"),
    ("estado", "Estado", "Expediente"),
    ("fecha_entrada", "Fecha de entrada", "Expediente"),
    ("observaciones", "Observaciones", "Expediente"),
    ("resumen", "Resumen del caso", "Expediente"),
    ("despachante", "Despachante asignado", "Expediente"),
    ("conexos", "Expedientes conexos", "Expediente"),
    # Del/los defendido/s (suele haber más de uno → numerados)
    ("defendidos", "Todos los defendidos (lista)", "Defendido"),
    ("defendido1", "1º defendido — nombre", "Defendido"),
    ("edad1", "1º defendido — edad", "Defendido"),
    ("dni1", "1º defendido — DNI", "Defendido"),
    ("defendido2", "2º defendido — nombre", "Defendido"),
    ("edad2", "2º defendido — edad", "Defendido"),
    ("dni2", "2º defendido — DNI", "Defendido"),
    ("defendido3", "3º defendido — nombre", "Defendido"),
    ("edad3", "3º defendido — edad", "Defendido"),
    ("dni3", "3º defendido — DNI", "Defendido"),
    ("defendido", "Defendido principal (= 1º)", "Defendido"),
    ("edad", "Edad del principal", "Defendido"),
    ("dni", "DNI del principal", "Defendido"),
    ("fecha_nacimiento", "Fecha de nacimiento (principal)", "Defendido"),
    ("vinculo", "Vínculo / rol (principal)", "Defendido"),
    # Institucionales / fecha
    ("defensora", "Nombre de la defensora", "Institucional"),
    ("defensoria", "Dependencia (Defensoría)", "Institucional"),
    ("ciudad", "Ciudad", "Institucional"),
    ("fecha", "Fecha de hoy (en letras)", "Institucional"),
]

# Alias que también se aceptan al escribir (apuntan a un token del catálogo).
ALIAS = {
    "expediente": "numero",
    "nro": "numero",
    "autos": "caratula",
    "hoy": "fecha",
}

ETIQUETAS = {tok: etq for tok, etq, _ in CATALOGO}

DEPENDENCIA = "Defensoría Pública de Menores e Incapaces N° 6"
CIUDAD = "Ciudad Autónoma de Buenos Aires"

_MESES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

# Variable: @ seguido de letras/números/guión bajo.
_TOKEN_RE = re.compile(r"@([a-zA-ZñÑáéíóúÁÉÍÓÚ_][a-zA-Z0-9ñÑáéíóúÁÉÍÓÚ_]*)")


def fecha_en_letras(d) -> str:
    if not d:
        return ""
    return f"{d.day} de {_MESES[d.month - 1]} de {d.year}"


def _edad(fnac, hoy) -> str:
    if not fnac:
        return ""
    años = hoy.year - fnac.year - ((hoy.month, hoy.day) < (fnac.month, fnac.day))
    return str(años)


def construir_contexto(db, exp) -> dict:
    """Arma el diccionario {token: valor} a partir del expediente y sus datos."""
    from app.models import Usuario

    hoy = date.today()
    defs = sorted(exp.defendidos, key=lambda d: d.id) if exp.defendidos else []
    d0 = defs[0] if defs else None
    defensora = db.query(Usuario).filter(Usuario.rol == "defensora").first()

    ctx = {
        "numero": exp.numero or "",
        "caratula": exp.caratula or "",
        "juzgado": exp.juzgado or "",
        "tipo_proceso": exp.tipo_proceso or "",
        "estado": exp.estado or "",
        "fecha_entrada": fecha_en_letras(exp.fecha_entrada),
        "observaciones": exp.observaciones or "",
        "resumen": exp.resumen or "",
        "despachante": exp.despachante_asignado or "",
        "conexos": ", ".join(exp.conexos) if exp.conexos else "",
        "defendido": (d0.nombre if d0 else "") or "",
        "defendidos": ", ".join(d.nombre for d in defs if d.nombre) if defs else "",
        "dni": (d0.dni if d0 else "") or "",
        "fecha_nacimiento": fecha_en_letras(d0.fecha_nacimiento) if d0 else "",
        "edad": _edad(d0.fecha_nacimiento, hoy) if d0 else "",
        "vinculo": (d0.vinculo if d0 else "") or "",
        "defensora": defensora.nombre if defensora else "",
        "defensoria": DEPENDENCIA,
        "ciudad": CIUDAD,
        "fecha": fecha_en_letras(hoy),
    }
    # Defendidos numerados (1..6): generalmente hay más de uno.
    for i in range(1, 7):
        d = defs[i - 1] if i - 1 < len(defs) else None
        ctx[f"defendido{i}"] = (d.nombre if d else "") or ""
        ctx[f"edad{i}"] = _edad(d.fecha_nacimiento, hoy) if d else ""
        ctx[f"dni{i}"] = (d.dni if d else "") or ""
        ctx[f"vinculo{i}"] = (d.vinculo if d else "") or ""
        ctx[f"nacimiento{i}"] = fecha_en_letras(d.fecha_nacimiento) if d else ""
    return ctx


def rellenar(texto: str, ctx: dict):
    """
    Reemplaza las @variables conocidas por su valor.
    - Variable conocida con dato     → el dato.
    - Variable conocida sin dato     → "[completar: Etiqueta]" (y se reporta).
    - Variable desconocida (ej. mail)→ se deja tal cual.
    Devuelve (texto_rellenado, lista_de_faltantes).
    """
    faltantes = []

    def _repl(m):
        token = m.group(1).lower()
        token = ALIAS.get(token, token)
        if token not in ctx:
            return m.group(0)
        valor = (ctx[token] or "").strip()
        if not valor:
            etq = ETIQUETAS.get(token, token)
            if etq not in faltantes:
                faltantes.append(etq)
            return f"[completar: {etq}]"
        return valor

    return _TOKEN_RE.sub(_repl, texto or ""), faltantes


def _rellenar_parrafo_docx(p, ctx, faltantes):
    """Reemplaza las @ en un párrafo de Word conservando el formato del párrafo."""
    if "@" not in p.text:
        return
    nuevo, falt = rellenar(p.text, ctx)
    for f in falt:
        if f not in faltantes:
            faltantes.append(f)
    if nuevo != p.text:
        if p.runs:
            p.runs[0].text = nuevo
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = nuevo


def rellenar_documento_docx(datos_docx: bytes, ctx: dict):
    """
    Rellena las @variables dentro de un .docx (párrafos y celdas de tablas),
    conservando el formato original. Devuelve (BytesIO del .docx, faltantes).
    """
    import io
    from docx import Document

    doc = Document(io.BytesIO(datos_docx))
    faltantes = []
    for p in doc.paragraphs:
        _rellenar_parrafo_docx(p, ctx, faltantes)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    _rellenar_parrafo_docx(p, ctx, faltantes)
    salida = io.BytesIO()
    doc.save(salida)
    salida.seek(0)
    return salida, faltantes
