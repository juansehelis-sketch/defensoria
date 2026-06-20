"""
Proyectos "a la firma": el flujo de trabajo interno que reemplaza el correo.

Flujo:
  1) Un remitente (despachante o secretaria) ENVÍA el proyecto de un expediente a
     un destinatario (secretaria o defensora), con archivos y datos.
     → se estampa "Pase a la firma" = hoy en el listado, y se notifica.
  2) El destinatario puede:
       - DEVOLVER con comentarios → estado "en_correccion", se notifica al remitente.
       - MARCAR SUBIDO → estado "subido", se estampa "Subido al Lex" = hoy, se notifica.
     La secretaria además puede REENVIAR A LA DEFENSORA.
  3) El remitente, si fue devuelto, REENVÍA una versión corregida.
Todo queda registrado en el historial del expediente.
"""

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from sqlalchemy.orm import Session
from sqlalchemy import or_
from datetime import datetime, date
from pathlib import Path
from typing import List
from html import escape
import shutil
import uuid

from app.database import get_db
from app.models import Proyecto, Expediente, EntradaSalida, Usuario, Notificacion, Historial
from app.schemas import Proyecto as ProyectoSchema
from app.utils.deps import obtener_usuario_actual
from app.services import storage

router = APIRouter(prefix="/api/proyectos", tags=["proyectos"])

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)


def _guardar_archivos(archivos: List[UploadFile]) -> list:
    """Guarda los PDFs adjuntos y devuelve [{nombre, url}]."""
    guardados = []
    for arch in archivos or []:
        if not arch or not arch.filename:
            continue
        ext = Path(arch.filename).suffix
        nombre_guardado = f"{uuid.uuid4().hex}{ext}"
        storage.guardar(nombre_guardado, arch.file.read(), arch.content_type)
        guardados.append({"nombre": arch.filename, "url": f"/uploads/{nombre_guardado}"})
    return guardados


def _entrada_para_estampar(db: Session, expediente: Expediente) -> EntradaSalida:
    """
    Devuelve la fila del listado sobre la que estampar las fechas del flujo.
    Usa la última entrada abierta (sin 'subido al lex'); si no hay, crea una nueva de hoy.
    """
    entrada = (
        db.query(EntradaSalida)
        .filter(EntradaSalida.expediente_id == expediente.id)
        .filter(EntradaSalida.subido_lex.is_(None))
        .order_by(EntradaSalida.fecha.desc())
        .first()
    )
    if entrada is None:
        entrada = EntradaSalida(
            fecha=date.today(),
            juzgado=expediente.juzgado,
            expediente_id=expediente.id,
            autos=expediente.caratula,
            asignacion=expediente.despachante.nombre if expediente.despachante else "",
            observaciones="",
            subido_defensa=False,
        )
        db.add(entrada)
        db.flush()
    return entrada


def _registrar_historial(db: Session, expediente_id: int, usuario: Usuario, tipo: str, desc: str):
    db.add(Historial(
        expediente_id=expediente_id,
        tipo=tipo,
        descripcion=desc,
        usuario_id=usuario.id,
    ))


def _docx_a_html(path: str) -> str:
    """Convierte un .docx a HTML simple (párrafos, negrita, itálica, títulos, tablas)."""
    from docx import Document
    doc = Document(path)
    partes = []
    for p in doc.paragraphs:
        texto = p.text
        if not texto.strip():
            partes.append("<br/>")
            continue
        estilo = (p.style.name if p.style else "") or ""
        tag = "p"
        if "Heading 1" in estilo or "Título 1" in estilo:
            tag = "h2"
        elif "Heading" in estilo or "Título" in estilo:
            tag = "h3"
        runs = ""
        for r in p.runs:
            t = escape(r.text)
            if r.bold:
                t = f"<strong>{t}</strong>"
            if r.italic:
                t = f"<em>{t}</em>"
            if r.underline:
                t = f"<u>{t}</u>"
            runs += t
        if not runs:
            runs = escape(texto)
        align = ""
        try:
            a = str(p.alignment) if p.alignment is not None else ""
            if "CENTER" in a:
                align = ' style="text-align:center"'
            elif "RIGHT" in a:
                align = ' style="text-align:right"'
            elif "JUSTIFY" in a:
                align = ' style="text-align:justify"'
        except Exception:
            pass
        partes.append(f"<{tag}{align}>{runs}</{tag}>")
    for table in doc.tables:
        partes.append('<table style="border-collapse:collapse;margin:8px 0">')
        for row in table.rows:
            celdas = "".join(
                f'<td style="border:1px solid #d4d9e2;padding:4px 8px">{escape(c.text)}</td>'
                for c in row.cells
            )
            partes.append(f"<tr>{celdas}</tr>")
        partes.append("</table>")
    return "\n".join(partes)


@router.get("/preview-docx")
async def preview_docx(url: str, usuario: Usuario = Depends(obtener_usuario_actual)):
    """Devuelve el contenido de un .docx como HTML para previsualizar sin descargar."""
    nombre = Path(url).name
    datos = storage.leer(nombre)
    if datos is None:
        raise HTTPException(status_code=404, detail="Archivo no encontrado")
    import tempfile, os
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    try:
        tmp.write(datos)
        tmp.close()
        return {"html": _docx_a_html(tmp.name)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo leer el documento: {e}")
    finally:
        os.unlink(tmp.name)


# ── Listados ───────────────────────────────────────────────────

@router.get("/recibidos", response_model=list[ProyectoSchema])
async def proyectos_recibidos(
    db: Session = Depends(get_db),
    usuario: Usuario = Depends(obtener_usuario_actual),
):
    """Proyectos que me enviaron (bandeja de entrada)."""
    return (
        db.query(Proyecto)
        .filter(Proyecto.destinatario_id == usuario.id)
        .order_by(Proyecto.fecha_envio.desc())
        .all()
    )


@router.get("/enviados", response_model=list[ProyectoSchema])
async def proyectos_enviados(
    db: Session = Depends(get_db),
    usuario: Usuario = Depends(obtener_usuario_actual),
):
    """Proyectos que yo envié."""
    return (
        db.query(Proyecto)
        .filter(Proyecto.remitente_id == usuario.id)
        .order_by(Proyecto.fecha_envio.desc())
        .all()
    )


@router.get("/expediente/{expediente_id}", response_model=list[ProyectoSchema])
async def proyectos_de_expediente(expediente_id: int, db: Session = Depends(get_db)):
    """Proyectos asociados a un expediente (se ven en la ficha)."""
    return (
        db.query(Proyecto)
        .filter(Proyecto.expediente_id == expediente_id)
        .order_by(Proyecto.fecha_envio.desc())
        .all()
    )


# ── Enviar un proyecto ─────────────────────────────────────────

@router.post("/", response_model=ProyectoSchema)
async def enviar_proyecto(
    expediente_id: int = Form(...),
    destinatario_id: int = Form(...),
    titulo: str = Form(""),
    datos: str = Form(""),
    archivos: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
    usuario: Usuario = Depends(obtener_usuario_actual),
):
    """Crea y envía un proyecto a la firma. Estampa 'Pase a la firma' = hoy."""
    expediente = db.query(Expediente).filter(Expediente.id == expediente_id).first()
    if not expediente:
        raise HTTPException(status_code=404, detail="Expediente no encontrado")

    destinatario = db.query(Usuario).filter(Usuario.id == destinatario_id).first()
    if not destinatario:
        raise HTTPException(status_code=404, detail="Destinatario no encontrado")
    if destinatario.rol not in ("secretaria", "defensora"):
        raise HTTPException(status_code=400, detail="Solo se puede enviar a una secretaria o a la defensora")

    guardados = _guardar_archivos(archivos)

    # Estampar pase a la firma en el listado
    entrada = _entrada_para_estampar(db, expediente)
    entrada.pase_firma = date.today()

    proyecto = Proyecto(
        expediente_id=expediente_id,
        entrada_salida_id=entrada.id,
        remitente_id=usuario.id,
        destinatario_id=destinatario_id,
        titulo=titulo or f"Proyecto {expediente.numero}",
        datos=datos,
        estado="enviado",
        version=1,
        archivos=guardados,
        comentarios=[{
            "autor": usuario.nombre, "rol": usuario.rol,
            "fecha": datetime.now().isoformat(), "texto": datos or "Proyecto enviado", "tipo": "envio",
        }],
    )
    db.add(proyecto)

    destino_txt = "la Defensora" if destinatario.rol == "defensora" else destinatario.nombre
    db.add(Notificacion(
        usuario_id=destinatario_id,
        tipo="proyecto_recibido",
        contenido=f"{usuario.nombre} envió a la firma de {destino_txt} — expte. {expediente.numero}",
        expediente_id=expediente_id,
    ))
    _registrar_historial(db, expediente_id, usuario, "otro",
                         f"Envió a la firma de {destino_txt}.")

    db.commit()
    db.refresh(proyecto)
    return proyecto


# ── Acciones del destinatario / remitente ──────────────────────

def _get_proyecto(db, proyecto_id) -> Proyecto:
    p = db.query(Proyecto).filter(Proyecto.id == proyecto_id).first()
    if not p:
        raise HTTPException(status_code=404, detail="Proyecto no encontrado")
    return p


@router.post("/{proyecto_id}/devolver", response_model=ProyectoSchema)
async def devolver_con_comentarios(
    proyecto_id: int,
    comentario: str = Form(...),
    db: Session = Depends(get_db),
    usuario: Usuario = Depends(obtener_usuario_actual),
):
    """El destinatario devuelve el proyecto con comentarios para corregir."""
    p = _get_proyecto(db, proyecto_id)
    if p.destinatario_id != usuario.id:
        raise HTTPException(status_code=403, detail="Solo el destinatario puede devolver el proyecto")

    p.estado = "en_correccion"
    p.comentarios = (p.comentarios or []) + [{
        "autor": usuario.nombre, "rol": usuario.rol,
        "fecha": datetime.now().isoformat(), "texto": comentario, "tipo": "devolucion",
    }]
    db.add(Notificacion(
        usuario_id=p.remitente_id,
        tipo="proyecto_devuelto",
        contenido=f"{usuario.nombre} lo devolvió con comentarios — expte. {p.expediente_numero}",
        expediente_id=p.expediente_id,
    ))
    _registrar_historial(db, p.expediente_id, usuario, "otro",
                         f"Proyecto devuelto con comentarios: {comentario}")
    from app.utils.auditoria import registrar
    registrar(db, usuario, "devolvió", "proyecto", f"Expte. {p.expediente_numero} — {p.titulo}")
    db.commit()
    db.refresh(p)
    return p


@router.post("/{proyecto_id}/reenviar", response_model=ProyectoSchema)
async def reenviar_corregido(
    proyecto_id: int,
    comentario: str = Form(""),
    archivos: List[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
    usuario: Usuario = Depends(obtener_usuario_actual),
):
    """El remitente reenvía una versión corregida tras una devolución."""
    p = _get_proyecto(db, proyecto_id)
    if p.remitente_id != usuario.id:
        raise HTTPException(status_code=403, detail="Solo el remitente puede reenviar el proyecto")

    guardados = _guardar_archivos(archivos)
    p.estado = "enviado"
    p.version = (p.version or 1) + 1
    p.fecha_envio = datetime.now()
    if guardados:
        p.archivos = (p.archivos or []) + guardados
    p.comentarios = (p.comentarios or []) + [{
        "autor": usuario.nombre, "rol": usuario.rol,
        "fecha": datetime.now().isoformat(),
        "texto": comentario or f"Versión corregida (v{p.version})", "tipo": "correccion",
    }]

    # Re-estampar pase a la firma con la fecha del reenvío
    if p.entrada_salida_id:
        entrada = db.query(EntradaSalida).filter(EntradaSalida.id == p.entrada_salida_id).first()
        if entrada:
            entrada.pase_firma = date.today()

    db.add(Notificacion(
        usuario_id=p.destinatario_id,
        tipo="proyecto_recibido",
        contenido=f"{usuario.nombre} reenvió una versión corregida del expte. {p.expediente_numero}",
        expediente_id=p.expediente_id,
    ))
    _registrar_historial(db, p.expediente_id, usuario, "otro",
                         f"Reenvió versión corregida (v{p.version}).")
    db.commit()
    db.refresh(p)
    return p


@router.post("/{proyecto_id}/reenviar-a-defensora", response_model=ProyectoSchema)
async def reenviar_a_defensora(
    proyecto_id: int,
    db: Session = Depends(get_db),
    usuario: Usuario = Depends(obtener_usuario_actual),
):
    """La secretaria reenvía el proyecto a la defensora."""
    if usuario.rol != "secretaria":
        raise HTTPException(status_code=403, detail="Solo una secretaria puede reenviar a la defensora")
    p = _get_proyecto(db, proyecto_id)
    if p.destinatario_id != usuario.id:
        raise HTTPException(status_code=403, detail="Solo podés reenviar un proyecto que recibiste")

    defensora = db.query(Usuario).filter(Usuario.rol == "defensora", Usuario.activo == True).first()
    if not defensora:
        raise HTTPException(status_code=400, detail="No hay una defensora cargada en el sistema")

    p.destinatario_id = defensora.id
    p.estado = "enviado"
    p.fecha_envio = datetime.now()
    p.comentarios = (p.comentarios or []) + [{
        "autor": usuario.nombre, "rol": usuario.rol,
        "fecha": datetime.now().isoformat(),
        "texto": f"Reenviado a la defensora ({defensora.nombre})", "tipo": "envio",
    }]
    db.add(Notificacion(
        usuario_id=defensora.id,
        tipo="proyecto_recibido",
        contenido=f"{usuario.nombre} envió a la firma de la Defensora — expte. {p.expediente_numero}",
        expediente_id=p.expediente_id,
    ))
    _registrar_historial(db, p.expediente_id, usuario, "otro", "Proyecto derivado a la defensora.")
    db.commit()
    db.refresh(p)
    return p


@router.post("/{proyecto_id}/subido", response_model=ProyectoSchema)
async def marcar_subido(
    proyecto_id: int,
    comentario: str = Form(""),
    dictamen: UploadFile = File(...),
    db: Session = Depends(get_db),
    usuario: Usuario = Depends(obtener_usuario_actual),
):
    """
    El destinatario confirma que subió el dictamen al expediente real.
    OBLIGA a adjuntar el PDF del dictamen subido: ese archivo queda guardado como
    dictamen del expediente (visible en el "mundo del expediente").
    Estampa 'Subido al Lex' = hoy en el listado y notifica al remitente.
    """
    p = _get_proyecto(db, proyecto_id)
    if p.destinatario_id != usuario.id:
        raise HTTPException(status_code=403, detail="Solo el destinatario puede marcarlo como subido")

    if not dictamen or not dictamen.filename:
        raise HTTPException(status_code=400, detail="Tenés que adjuntar el PDF del dictamen subido")

    # Guardar el dictamen final
    guardados = _guardar_archivos([dictamen])
    dictamen_url = guardados[0]["url"] if guardados else None

    p.estado = "subido"
    p.fecha_subido = datetime.now()
    if guardados:
        p.archivos = (p.archivos or []) + guardados
    p.comentarios = (p.comentarios or []) + [{
        "autor": usuario.nombre, "rol": usuario.rol,
        "fecha": datetime.now().isoformat(),
        "texto": comentario or "Subió el dictamen al expediente", "tipo": "subido",
    }]

    # Estampar subido al lex en el listado
    if p.entrada_salida_id:
        entrada = db.query(EntradaSalida).filter(EntradaSalida.id == p.entrada_salida_id).first()
        if entrada:
            entrada.subido_lex = date.today()
            entrada.subido_defensa = True

    db.add(Notificacion(
        usuario_id=p.remitente_id,
        tipo="proyecto_subido",
        contenido=f"{usuario.nombre} subió el dictamen al expediente — expte. {p.expediente_numero}",
        expediente_id=p.expediente_id,
    ))
    # Guardar el dictamen como intervención del expediente (queda en su "mundo")
    db.add(Historial(
        expediente_id=p.expediente_id,
        tipo="dictamen",
        descripcion=f"Dictamen subido al expediente por {usuario.nombre}." + (f" {comentario}" if comentario else ""),
        usuario_id=usuario.id,
        archivo_url=dictamen_url,
    ))
    from app.utils.auditoria import registrar
    registrar(db, usuario, "subió", "proyecto", f"Dictamen subido — expte. {p.expediente_numero}")
    db.commit()
    db.refresh(p)
    return p
