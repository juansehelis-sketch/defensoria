"""
Endpoints de audiencias: ABM, vista por rango de fechas (calendario)
e importación desde texto pegado (formato de la planilla).
"""

from fastapi import APIRouter, Depends, HTTPException, status, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from datetime import date, time, datetime
import io
from app.database import get_db
from app.models import Audiencia, Expediente
from app.schemas import Audiencia as AudienciaSchema, AudienciaCreate
from app.utils.deps import obtener_usuario_actual
from app.models import Usuario

router = APIRouter(prefix="/api/audiencias", tags=["audiencias"])


@router.get("/", response_model=list[AudienciaSchema])
async def listar_audiencias(
    fecha_inicio: date = Query(None),
    fecha_fin: date = Query(None),
    db: Session = Depends(get_db),
):
    """Lista audiencias, opcionalmente filtradas por rango de fechas (para el calendario)."""
    query = db.query(Audiencia)
    if fecha_inicio:
        query = query.filter(Audiencia.fecha >= fecha_inicio)
    if fecha_fin:
        query = query.filter(Audiencia.fecha <= fecha_fin)
    return query.order_by(Audiencia.fecha.asc(), Audiencia.hora.asc()).all()


@router.get("/mias", response_model=list[AudienciaSchema])
async def mis_audiencias(
    db: Session = Depends(get_db),
    usuario: Usuario = Depends(obtener_usuario_actual),
):
    """
    Agenda personal: solo las audiencias asignadas al usuario logueado
    (las que tienen su nombre en 'quién va'). Ordenadas por fecha y hora.
    """
    return (
        db.query(Audiencia)
        .filter(Audiencia.asignado_a == usuario.nombre)
        .order_by(Audiencia.fecha.asc(), Audiencia.hora.asc())
        .all()
    )


@router.post("/", response_model=AudienciaSchema)
async def crear_audiencia(audiencia: AudienciaCreate, db: Session = Depends(get_db)):
    """Crea una audiencia vinculada a un expediente."""
    expediente = db.query(Expediente).filter(Expediente.id == audiencia.expediente_id).first()
    if not expediente:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Expediente no encontrado")

    nueva = Audiencia(
        expediente_id=audiencia.expediente_id,
        fecha=audiencia.fecha,
        hora=audiencia.hora,
        juzgado=audiencia.juzgado,
        base_legal=audiencia.base_legal,
        motivo=audiencia.motivo,
        modalidad=audiencia.modalidad,
        datos_acceso=audiencia.datos_acceso,
        direccion=audiencia.direccion,
        asesor=audiencia.asesor,
        asignado_a=audiencia.asignado_a,
        asistencia=audiencia.asistencia or "pendiente",
        estado=audiencia.estado,
    )
    db.add(nueva)
    db.commit()
    db.refresh(nueva)
    return nueva


@router.put("/{audiencia_id}", response_model=AudienciaSchema)
async def actualizar_audiencia(audiencia_id: int, datos: dict, db: Session = Depends(get_db)):
    """Actualiza una audiencia (fecha, hora, estado, etc.)."""
    audiencia = db.query(Audiencia).filter(Audiencia.id == audiencia_id).first()
    if not audiencia:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Audiencia no encontrada")
    for key, value in datos.items():
        if hasattr(audiencia, key) and value is not None:
            setattr(audiencia, key, value)
    db.commit()
    db.refresh(audiencia)
    return audiencia


@router.delete("/{audiencia_id}")
async def eliminar_audiencia(audiencia_id: int, db: Session = Depends(get_db)):
    """Elimina una audiencia."""
    audiencia = db.query(Audiencia).filter(Audiencia.id == audiencia_id).first()
    if not audiencia:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Audiencia no encontrada")
    db.delete(audiencia)
    db.commit()
    return {"message": "Audiencia eliminada"}


# ── Acta de audiencia (Word para completar y firmar) ───────────

@router.post("/{audiencia_id}/acta")
async def generar_acta(audiencia_id: int, db: Session = Depends(get_db)):
    """
    Genera un acta de audiencia en Word (.docx) ya prellenada con los datos
    (fecha, hora, juzgado, expediente, carátula) y con espacio para completar
    lo ocurrido. Se descarga y se termina en Word.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from app.services import plantillas as plantillas_svc

    a = db.query(Audiencia).filter(Audiencia.id == audiencia_id).first()
    if not a:
        raise HTTPException(status_code=404, detail="Audiencia no encontrada")

    exp = a.expediente
    numero = (exp.numero if exp else None) or "—"
    caratula = (exp.caratula if exp else None) or "—"
    juzgado = a.juzgado or (exp.juzgado if exp else None) or "—"
    fecha_letras = plantillas_svc.fecha_en_letras(a.fecha) if a.fecha else "—"
    hora = a.hora.strftime("%H:%M") if a.hora else "—"

    doc = Document()
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("ACTA DE AUDIENCIA")
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph()

    def campo(label, valor):
        p = doc.add_paragraph()
        b = p.add_run(f"{label}: ")
        b.bold = True
        p.add_run(valor or "—")

    campo("Fecha", fecha_letras)
    campo("Hora", hora)
    campo("Juzgado", juzgado)
    campo("Expediente N°", numero)
    campo("Autos", caratula)
    if a.modalidad:
        campo("Modalidad", a.modalidad)
    campo("Por la Defensoría", a.asignado_a or "")
    if a.motivo:
        campo("Motivo", a.motivo)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.add_run(
        f"En la audiencia celebrada el día {fecha_letras}, siendo las {hora} horas, "
        f"en los autos caratulados \"{caratula}\" (Expte. N° {numero}) que tramitan ante el "
        f"{juzgado}, comparece la Defensoría Pública de Menores e Incapaces N° 6, y se deja "
        "constancia de lo siguiente:"
    )
    for _ in range(8):
        doc.add_paragraph()

    cierre = doc.add_paragraph()
    cierre.add_run(
        "No siendo para más, se da por finalizado el acto, previa lectura y ratificación de los comparecientes."
    )
    doc.add_paragraph()
    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma.add_run("_______________________________")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    base = "".join(c for c in numero if c.isalnum() or c in "-_") or "audiencia"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="acta_{base}.docx"'},
    )


# ── Importación desde texto pegado ─────────────────────────────

class TextoImportar(BaseModel):
    texto: str


def _parse_fecha(s: str):
    """Convierte 'dd/mm/aaaa' (o con guiones) a date."""
    s = s.strip().replace("-", "/")
    partes = s.split("/")
    if len(partes) < 3:
        return None
    try:
        return date(int(partes[2]), int(partes[1]), int(partes[0]))
    except (ValueError, IndexError):
        return None


def _parse_hora(s: str):
    """Convierte 'HH:MM' (o 'HH') a time."""
    s = s.strip().lower().replace("hs", "").replace(".", "").strip()
    try:
        if ":" in s:
            h, m = s.split(":")[:2]
            return time(int(h), int(m))
        return time(int(s), 0)
    except (ValueError, IndexError):
        return None


@router.post("/importar-texto")
async def importar_desde_texto(
    datos: TextoImportar,
    db: Session = Depends(get_db),
    usuario: Usuario = Depends(obtener_usuario_actual),
):
    """
    Importa audiencias desde texto pegado de la planilla.
    Columnas esperadas (separadas por tabulación):
    fecha | hora | juzgado | expediente | caratula | base legal | asesor | modalidad

    Vincula cada audiencia al expediente por su número (si existe).
    Devuelve un preview con creadas y advertencias (no requiere confirmación previa).
    """
    lineas = [l for l in datos.texto.split("\n") if l.strip()]
    creadas = []
    advertencias = []

    for i, linea in enumerate(lineas, 1):
        cols = [c.strip() for c in linea.split("\t")]
        if len(cols) < 8:
            advertencias.append(f"Fila {i} ignorada (menos de 8 columnas)")
            continue

        fecha = _parse_fecha(cols[0])
        if not fecha:
            advertencias.append(f"Fila {i} ignorada (fecha no reconocida: '{cols[0]}')")
            continue

        hora = _parse_hora(cols[1]) or time(0, 0)
        juzgado = cols[2]
        num_expte = cols[3]
        base_legal = cols[5]
        asesor = cols[6]
        modalidad = " ".join(cols[7:]).strip()

        # Buscar expediente por número (normalizado simple)
        expediente = db.query(Expediente).filter(Expediente.numero == num_expte).first()
        if not expediente:
            # Buscar quitando el asterisco final típico de la planilla
            expediente = db.query(Expediente).filter(
                Expediente.numero == num_expte.replace("*", "").strip()
            ).first()

        if not expediente:
            advertencias.append(
                f"Fila {i}: expediente '{num_expte}' no existe en el sistema; audiencia no creada"
            )
            continue

        nueva = Audiencia(
            expediente_id=expediente.id,
            fecha=fecha,
            hora=hora,
            juzgado=juzgado,
            base_legal=base_legal,
            modalidad=modalidad,
            asesor=asesor,
            estado="programada",
        )
        db.add(nueva)
        db.commit()
        db.refresh(nueva)
        creadas.append({"id": nueva.id, "expediente": num_expte, "fecha": str(fecha)})

    return {
        "creadas": creadas,
        "total_creadas": len(creadas),
        "advertencias": advertencias,
    }
